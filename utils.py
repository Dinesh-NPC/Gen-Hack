import os
import tempfile
from typing import List, Tuple
from PIL import Image
from pypdf import PdfReader
from docx import Document
from langchain_core.documents import Document as LangchainDocument
from sentence_transformers import SentenceTransformer
from transformers import BlipProcessor, BlipForConditionalGeneration
import torch

# Initialize models
try:
    import whisper
    whisper_model = whisper.load_model("base")
except Exception as e:
    print(f"Warning: Whisper import failed: {e}. Audio processing will be disabled.")
    whisper = None
    whisper_model = None

try:
    clip_model = SentenceTransformer('clip-ViT-B-32')
    blip_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
    blip_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
except Exception as e:
    print(f"Warning: Model loading failed: {e}. Some features may be disabled.")
    clip_model = None
    blip_processor = None
    blip_model = None

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file."""
    with open(file_path, 'rb') as file:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def transcribe_audio(file_path: str) -> str:
    """Transcribe audio file to text using Whisper."""
    result = whisper_model.transcribe(file_path)
    return result["text"]

def generate_image_caption(image_path: str) -> str:
    """Generate caption for image using BLIP."""
    image = Image.open(image_path).convert('RGB')
    inputs = blip_processor(image, return_tensors="pt")
    out = blip_model.generate(**inputs)
    caption = blip_processor.decode(out[0], skip_special_tokens=True)
    return caption

def get_image_embedding(image_path: str) -> List[float]:
    """Get CLIP embedding for image."""
    image = Image.open(image_path).convert('RGB')
    embedding = clip_model.encode(image)
    return embedding.tolist()

def get_text_embedding(text: str) -> List[float]:
    """Get CLIP embedding for text (for multimodal consistency)."""
    embedding = clip_model.encode(text)
    return embedding.tolist()

def process_file(file_path: str, file_type: str) -> Tuple[str, List[float], str]:
    """
    Process a file and return content, embedding, and metadata.
    Returns: (content, embedding, metadata_type)
    """
    if file_type == 'text':
        if file_path.endswith('.pdf'):
            content = extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            content = extract_text_from_docx(file_path)
        elif file_path.endswith('.txt'):
            content = extract_text_from_txt(file_path)
        else:
            content = ""
        embedding = get_text_embedding(content)
        metadata_type = 'text'
    elif file_type == 'image':
        content = generate_image_caption(file_path)
        embedding = get_image_embedding(file_path)
        metadata_type = 'image'
    elif file_type == 'audio':
        content = transcribe_audio(file_path)
        embedding = get_text_embedding(content)
        metadata_type = 'audio'
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    return content, embedding, metadata_type

def create_langchain_documents(contents: List[str], metadatas: List[dict]) -> List[LangchainDocument]:
    """Create LangChain Document objects from contents and metadatas."""
    documents = []
    for content, metadata in zip(contents, metadatas):
        doc = LangchainDocument(page_content=content, metadata=metadata)
        documents.append(doc)
    return documents
